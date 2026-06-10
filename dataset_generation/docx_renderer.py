
import uuid
import shutil

from docx import Document

from config_ import FIELDS

class DocxRenderer:

    def __init__(self, template_path):

        self.template_path = template_path


    def create_document(

        self,

        output_path,

        sample

    ):

        shutil.copy(

            self.template_path,

            output_path

        )

        doc = Document(

            output_path

        )

        marker_map = {}

        value_map = {}

        for paragraph in doc.paragraphs:

            self._replace_in_paragraph(

                paragraph,

                sample,

                marker_map,

                value_map

            )

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        self._replace_in_paragraph(

                            paragraph,

                            sample,

                            marker_map,

                            value_map

                        )

        doc.save(

            output_path

        )

        return marker_map, value_map


    def _replace_in_paragraph(

        self,

        paragraph,

        sample,

        marker_map,

        value_map

    ):

        text = paragraph.text

        changed = False

        for placeholder, field in FIELDS.items():

            if placeholder in text:

                marker = make_marker(field)

                text = text.replace(

                    placeholder,

                    marker

                )

                marker_map[marker] = field

                value_map[marker] = sample[field]

                changed = True

        if not changed:

            return

        for run in paragraph.runs:

            run.text = ""

        if len(paragraph.runs) == 0:

            paragraph.add_run(text)

        else:

            paragraph.runs[0].text = text


    def _create_marker(

        self,

        field

    ):

        return (

            "<<"

            + field.upper()

            + "_"

            + uuid.uuid4().hex

            + ">>"

        )
def make_marker(field: str, index: int = 1):

    prefix={

        "patient_name": "P",

        "birth_date": "B",

        "doctor_name": "D",

        "medicine": "M",

        "form": "F",

        "dosage": "G",

        "days": "Y"

    }[field]

    return f"§{prefix}{index:03d}§"
